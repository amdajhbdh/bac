import os
import sys
import json
import uuid
import hashlib
import PyPDF2
import docx
import psycopg2
from psycopg2.extras import Json
import numpy as np
import requests

# Vector DB configuration (Neon/Postgres)
DB_URL = os.getenv("NEON_DB_URL")
# Local Ollama endpoint for embeddings
EMBED_ENDPOINT = "http://localhost:11434/api/embeddings"
EMBED_MODEL = "nomic-embed-text"

class KBProcessor:
    def __init__(self, db_url):
        self.conn = psycopg2.connect(db_url)
        self.cursor = self.conn.cursor()

    def get_embedding(self, text):
        """Generate vector embedding using local Ollama (Llama 3 family)"""
        try:
            response = requests.post(EMBED_ENDPOINT, json={
                "model": EMBED_MODEL,
                "prompt": text
            })
            return response.json()["embedding"]
        except Exception as e:
            print(f"Error generating embedding: {e}")
            return [0.0] * 1536  # Fallback zero-vector

    def extract_pdf_metadata(self, file_path):
        """Extract text and metadata from PDF"""
        text = ""
        metadata = {}
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            metadata = {
                "pages": len(reader.pages),
                "author": reader.metadata.get('/Author', ''),
                "creator": reader.metadata.get('/Creator', ''),
                "producer": reader.metadata.get('/Producer', ''),
                "subject": reader.metadata.get('/Subject', '')
            }
            # Extract first 5000 chars for embedding/content
            for page in reader.pages[:10]:
                text += page.extract_text() + "
"
        
        return text, metadata

    def extract_docx_metadata(self, file_path):
        """Extract text and metadata from DOCX"""
        doc = docx.Document(file_path)
        text = "
".join([p.text for p in doc.paragraphs])
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "author": doc.core_properties.author,
            "created": str(doc.core_properties.created)
        }
        return text, metadata

    def process_file(self, file_path, subject="General"):
        """Process file and insert into vector KB"""
        ext = os.path.splitext(file_path)[1].lower()
        title = os.path.basename(file_path)
        
        print(f"Processing: {title} ({ext})")
        
        try:
            if ext == ".pdf":
                content, metadata = self.extract_pdf_metadata(file_path)
            elif ext == ".docx":
                content, metadata = self.extract_docx_metadata(file_path)
            elif ext == ".txt":
                with open(file_path, 'r') as f:
                    content = f.read()
                metadata = {"size": os.path.getsize(file_path)}
            else:
                return

            # Clean content for embedding
            clean_content = content[:4000] # Cap for embedding efficiency
            embedding = self.get_embedding(clean_content)
            
            # Generate summary for token management (optional)
            summary = self.summarize(clean_content)

            # Insert into Postgres knowledge_base table
            self.cursor.execute("""
                INSERT INTO knowledge_base (
                    id, title, content, metadata, subject, category, 
                    source_type, source_path, embedding, summary
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                ON CONFLICT (id) DO UPDATE SET
                    embedding = EXCLUDED.embedding,
                    updated_at = NOW()
            """, (
                str(uuid.uuid4()), 
                title, 
                content, 
                Json(metadata), 
                subject, 
                "BAC", 
                ext[1:], 
                file_path, 
                embedding,
                summary
            ))
            self.conn.commit()
            print(f"✓ Saved: {title}")

        except Exception as e:
            print(f"Error processing {title}: {e}")

    def summarize(self, text):
        """Summarize content to reduce token count for NotebookLM"""
        # Placeholder for LLM-based summarization logic
        return text[:1000] + "..." if len(text) > 1000 else text

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python process_kb.py <file_path> [subject]")
        sys.exit(1)
    
    processor = KBProcessor(DB_URL)
    subject = sys.argv[2] if len(sys.argv) > 2 else "General"
    processor.process_file(sys.argv[1], subject)
